"""多格式文档解析器。

支持：
- .txt / .md  : 直接读取
- .html / .htm: BeautifulSoup 去除 nav/script/style/标签
- .docx       : python-docx
- .pdf        : pypdf（足够稳定，无需额外系统依赖）

TODO(OCR): 后续可在此处接入 PaddleOCR / RapidOCR 处理扫描版 PDF 与图片。
"""

from __future__ import annotations

from pathlib import Path

from app.utils.logger import logger


class DocumentLoadError(Exception):
    """把“文件不存在 / 格式不支持 / 解析失败”与业务 HTTP 400 对齐的专用异常。"""


def load_document(file_path: str | Path) -> str:
    """对外的统一门面（Facade）函数：你只需要传路径，内部根据扩展名跳转。

    为什么要分函数而不是一个巨大 if？
    ——可读性更好，也方便单元测试对每个格式单独断言。
    """
    path = Path(file_path)
    if not path.exists():
        raise DocumentLoadError(f"file not found: {path}")

    suffix = path.suffix.lower()
    logger.info(f"loading document: {path.name} ({suffix})")

    if suffix in {".txt", ".md"}:
        return _load_text(path)
    if suffix in {".html", ".htm"}:
        return _load_html(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix == ".pdf":
        return _load_pdf(path)

    raise DocumentLoadError(f"unsupported file type: {suffix}")


# ---------------- 内部实现 ----------------

def _load_text(path: Path) -> str:
    """txt / md 容错读取，自动尝试常见编码。"""
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_html(path: Path) -> str:
    """HTML → 近似“阅读模式”正文。

    BeautifulSoup (`lxml` 解析器) 能快速去掉标签层级；顺带剔除导航栏、脚本，减少噪声。
    """
    from bs4 import BeautifulSoup

    raw = _load_text(path)
    soup = BeautifulSoup(raw, "lxml")
    for tag in soup(["script", "style", "noscript", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    return text


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_pdf(path: Path) -> str:
    """PDF 文本层抽取（适用于“可选择复制”的电子 PDF）。

    如果是扫描拍照件，这里没有像素 → OCR，`extract_text()` 通常会得到空字符串。
    此时应在业务层提示用户换文本版或接入 OCR（见下方 TODO 注释）。
    """
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:  # 单页失败不要让整篇文档失败
            logger.warning(f"pdf page {i} extract failed: {e}")
            text = ""
        if text.strip():
            pages.append(text)
    if not pages:
        # TODO(OCR): 这里可触发 OCR fallback
        logger.warning(
            f"pdf {path.name} 文本提取为空，可能是扫描版 PDF，建议接入 OCR。"
        )
    return "\n".join(pages)
