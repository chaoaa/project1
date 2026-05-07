"""业务规则工具集：确定性规则判断 + 文档导出 + 用户画像抽取。

=============================================================================
设计原则（为什么要有这个文件？）
=============================================================================
1. 资格判断这类高风险决策，必须用确定性规则而非 LLM 猜测。
   LLM 可能产生幻觉（hallucination）——明明有挂科却说"可以申请"。
   规则引擎的结果是 100% 可复现、可审计的。

2. LLM 只负责两件事：
   - "理解用户自然语言" → 抽取结构化字段（但我们这里连抽取也用规则做了，更可控）
   - "解释规则结果" → 把规则引擎的判断结果用自然语言包装得更好看

3. 所有规则函数都是纯 Python，零外部 API 调用，可单测、可审计。
   你可以在 pytest 里直接 import 并测试，不需要启动 LLM。

为什么单独一个文件？
- 与 tools.py（LLM 驱动的工具）职责不同：这里全是确定性逻辑。
- 方便后续为不同学校定制规则而不影响 Agent 编排层。
  例如：北大和清华的奖学金规则不同 → 只需改这个文件。
=============================================================================
"""

from __future__ import annotations

import re
from datetime import datetime
from typing import Any, Dict, List, Optional

from app.config import settings
from app.utils.logger import logger


# =========================================================
# 功能 1：规则化奖学金资格判断
# =========================================================
# 这部分实现了"从用户自然语言中抽取结构化画像 → 确定性规则引擎判断"的完整链路。
# 和纯 LLM 方案的区别：
#   LLM 方案：问题 + 政策 → LLM → "你符合条件"（可能胡说）
#   本方案：  问题 → 正则抽取画像 → 规则引擎（纯 Python if/else）→ 确定结论
#           政策 → LLM → 补充依据引用（LLM 只负责"引用原文"这个安全任务）


def extract_scholarship_profile_from_question(
    question: str,
    user_profile: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """从用户自然语言问题中抽取结构化的学生画像字段。

    【核心策略】
    正则 + 关键词规则（确定性），完全不依赖大模型。
    原因：大模型抽取虽然更"智能"，但可能：
    - 把"没有挂科"误抽成 has_failed_course=True（否定句理解错误）
    - 对数字不敏感（"前20%"可能抽成 0.2 或 20，不确定）
    用正则虽然覆盖不全，但每一条都是可解释、可调试的。

    【优先级规则】
    user_profile（用户显式传入的 JSON）> question（从自然语言抽取）。
    因为用户在前端表单填的 JSON 比口语表达更精确。

    【支持的字段及含义】
    - grade:           年级，如"研一""研二"。"研二"和"硕士二年级"归一化为"研二"。
    - rank_percent:    成绩排名百分比（float）。"前20%" → 20.0，"排名前 15%" → 15.0。
    - failed_courses:  挂科/不及格课程数量（int）。
    - has_failed_course: 是否有挂科（bool）。注意区分"没有挂科"（否定）和"有一门不及格"（肯定）。
    - has_discipline:  是否有处分记录。
    - has_academic_misconduct: 是否有学术不端行为。
    - has_fake_material: 是否材料造假。
    - tutor_negative:  导师评价是否不合格。

    【扩展指南】
    想支持更多字段？在这个函数里加正则规则即可，不影响下游。
    """

    # ---- 初始化：所有字段给默认值，避免下游 KeyError ----
    profile: Dict[str, Any] = {
        "grade": None,              # None = 未识别到
        "rank_percent": None,       # None = 未识别到
        "failed_courses": 0,        # 默认 0 门
        "has_failed_course": False, # 默认没有
        "has_discipline": False,
        "has_academic_misconduct": False,
        "has_fake_material": False,
        "tutor_negative": False,
    }

    # ---- 第 1 步：优先从 user_profile 提取已知字段 ----
    # user_profile 是前端表单传入的 JSON，例如 {"年级":"研二","成绩排名":"前20%","挂科情况":"有1门不及格"}
    if user_profile:
        _merge_user_profile(profile, user_profile)

    # ---- 第 2 步：从 question 的自然语言中抽取补充字段 ----
    # 只填充 profile 中仍为 None/False/0 的字段（user_profile 已有值的不会被覆盖）
    text = question or ""

    # 2a. 抽取年级
    # 技巧：用映射表把各种表述归一化。例如"硕士二年级" → "研二"，方便后续规则匹配。
    if profile["grade"] is None:
        grade_map = {
            "研一": "研一", "研二": "研二", "研三": "研三",
            "博一": "博一", "博二": "博二", "博三": "博三", "博四": "博四",
            "大一": "大一", "大二": "大二", "大三": "大三", "大四": "大四",
            "硕士一年级": "研一", "硕士二年级": "研二", "硕士三年级": "研三",
            "博士一年级": "博一", "博士二年级": "博二", "博士三年级": "博三",
        }
        for key, val in grade_map.items():
            if key in text:
                profile["grade"] = val
                break  # 找到第一个就停，避免"研一"和"研二"同时命中

    # 2b. 抽取排名百分比
    # 正则解读：
    #   r"排名[前]?\s*(\d+)\s*%"  → 匹配 "排名前20%"、"排名前 20%"、"排名20%"
    #   [前]? 表示"前"字可选（口语中可能省略）
    #   (\d+) 捕获数字部分
    if profile["rank_percent"] is None:
        rank_match = re.search(r"排名[前]?\s*(\d+)\s*%", text)
        if rank_match:
            profile["rank_percent"] = float(rank_match.group(1))
        else:
            # 降级匹配：没有"排名"关键词，只有"前X%"
            rank_match2 = re.search(r"[前]?\s*(\d+)\s*%", text)
            if rank_match2:
                val = float(rank_match2.group(1))
                # 防守：百分比应在 0-100 之间，避免把"2024年"的"20"误识别
                if val <= 100:
                    profile["rank_percent"] = val

    # 2c. 抽取挂科信息 — 正向（有挂科）
    # 注意：这里要先判断"有X门不及格"，再判断"挂科"关键词。
    # 因为"有一门不及格"比"挂科"携带了更多信息（具体数量）。
    if not profile["has_failed_course"] and profile["failed_courses"] == 0:
        # 模式 1：精确数量 + 挂科/不及格
        # 例："有1门不及格"、"有一门课程挂科"、"有 2 门挂科"
        if re.search(r"有\s*(一|二|三|四|五|1|2|3|4|5)\s*[门科]\s*(挂|不及格)", text):
            profile["has_failed_course"] = True
            # 尝试提取具体数字
            num_match = re.search(r"(\d+)\s*[门科]", text)
            if num_match:
                profile["failed_courses"] = int(num_match.group(1))
            else:
                # 中文数字映射
                cn_num = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5}
                cn_match = re.search(r"([一二三四五])\s*[门科]", text)
                if cn_match:
                    profile["failed_courses"] = cn_num.get(cn_match.group(1), 1)
                else:
                    profile["failed_courses"] = 1  # 默认 1 门

        # 模式 2：模糊表述，没有精确数量
        # 例："挂科了"、"有不及格"、"挂了一科"
        # 但必须排除否定句："没有挂科"不应匹配
        elif (
            re.search(r"(挂科|不及格|挂了一科|挂了.*科|一门.*不及格)", text)
            and not re.search(r"(没有|无|未|没)\s*(挂科|不及格)", text)
        ):
            profile["has_failed_course"] = True
            profile["failed_courses"] = 1

    # 2d. 抽取挂科信息 — 否定（没有挂科）
    # 例："没有挂科"、"无不及格"、"未挂科"
    # 但需要小心："我没有挂科，但是有一门不及格"中"没有挂科"不意味着真的没有
    # 所以加上 _has_positive_failed 二次确认
    if re.search(r"(没有|无|没|未)\s*(挂科|不及格|重修)", text):
        if not _has_positive_failed(text):
            profile["has_failed_course"] = False
            profile["failed_courses"] = 0

    # 2e. 抽取处分信息
    # 正向："有处分"、"受过处分"、"记过处分"、"处分未解除"
    # 否定："没有处分"、"无处分"
    if re.search(r"(有|受过|被|记过|警告|留校察看|处分未解除)", text) and re.search(r"处分", text):
        if not re.search(r"(没有|无|未|没)\s*处分", text):
            profile["has_discipline"] = True
    if re.search(r"(没有|无|未|没)\s*处分", text):
        profile["has_discipline"] = False

    # 2f. 抽取学术不端
    # 注意：这个字段很少在口语中出现，但如果出现就是严重问题
    if re.search(r"学术不端", text):
        if not re.search(r"(没有|无|未|没)\s*学术不端", text):
            profile["has_academic_misconduct"] = True

    # 2g. 抽取材料造假
    if re.search(r"材料造假", text):
        if not re.search(r"(没有|无|未|没)\s*材料造假", text):
            profile["has_fake_material"] = True

    # 2h. 抽取导师评价
    if re.search(r"导师评价.*(不合格|差|不过|不通过)", text) or re.search(r"导师.*不合格", text):
        profile["tutor_negative"] = True

    logger.info(f"[business_tools] extracted profile: {profile}")
    return profile


def _merge_user_profile(profile: Dict[str, Any], user_profile: Dict[str, Any]) -> None:
    """将前端传入的 user_profile JSON 合并到 profile 字典。

    做简单的归一化处理：
    - "前20%"、"20%"、"20" 统一解析为 float 20.0
    - "有1门不及格" → has_failed_course=True, failed_courses=1
    - "是"/"true"/"1" → has_discipline=True
    """
    # 年级：直接使用
    if user_profile.get("年级") and profile["grade"] is None:
        profile["grade"] = str(user_profile["年级"])

    # 成绩排名：支持 "前20%"、"20%"、"20" 等多种写法
    if user_profile.get("成绩排名") and profile["rank_percent"] is None:
        rank_str = str(user_profile["成绩排名"])
        rank_match = re.search(r"(\d+\.?\d*)", rank_str)
        if rank_match:
            profile["rank_percent"] = float(rank_match.group(1))

    # 挂科情况：解析"有1门不及格"、"有挂科"等
    if user_profile.get("挂科情况") and not profile["has_failed_course"]:
        val = str(user_profile["挂科情况"])
        if re.search(r"(有|挂|不及格|\d+)\s*[门科]?", val) and not re.search(r"(没有|无|未|没)", val):
            profile["has_failed_course"] = True
            num_match = re.search(r"(\d+)", val)
            profile["failed_courses"] = int(num_match.group(1)) if num_match else 1

    # 挂科数量：直接数字
    if user_profile.get("挂科数量") is not None:
        try:
            profile["failed_courses"] = int(user_profile["挂科数量"])
            if profile["failed_courses"] > 0:
                profile["has_failed_course"] = True
        except (ValueError, TypeError):
            pass  # 无法转为 int 就忽略，不抛异常

    # 处分：支持 boolean 和中文字符串
    if user_profile.get("处分") is not None and not profile["has_discipline"]:
        val = str(user_profile["处分"])
        if val.lower() in ("有", "是", "true", "yes", "1"):
            profile["has_discipline"] = True


def _has_positive_failed(text: str) -> bool:
    """二次确认：检查文本中是否真的存在正向的挂科描述。

    为什么需要这个函数？
    例："我没有挂科"中出现了"没有"+"挂科"，但这不代表有挂科。
    但如果文本是"我没有挂科，但是有一门不及格"，则需要识别出确实有不及格。

    本函数只检查"有X门不及格/挂科"模式，不检查否定句。
    配合调用方的逻辑：先检查否定（"没有挂科"），再用本函数确认是否有正向描述。
    """
    return bool(re.search(
        r"(有\s*(一|二|三|四|五|1|2|3|4|5)?\s*[门科]?\s*(挂|不及格|重修|没过))|"
        r"(挂科|不及格|挂了一科|挂了.*科|一门.*不及格)",
        text,
    ))


def rule_based_scholarship_checker(profile: Dict[str, Any]) -> Dict[str, Any]:
    """确定性规则奖学金资格判断工具（纯 Python if/else，不调用任何外部 API）。

    【为什么这个工具必须是确定性的？】
    资格判断直接影响学生能否申请奖学金。如果用 LLM 直接判断：
    - 同一个人问两次可能得到不同答案（temperature > 0）
    - LLM 可能忽略关键的一票否决条件（如挂科）
    - 无法给教务处解释"为什么这个学生被判定为不合格"

    因此这里的规则引擎是纯 Python 逻辑：相同输入永远得到相同输出。

    【规则来源】
    2025 版研究生学业奖学金政策（示例）。实际部署时需要根据学校真实政策调整。

    【判断流程】
    1. 一票否决项检查（任意一项命中 → 直接 not_eligible）
       - 有挂科/不及格
       - 有处分记录
       - 有学术不端
       - 材料造假
       - 导师评价不合格
    2. 排名判断（需要 rank_percent 字段不为 None）
       - 前 20%（含）→ 一等学业奖学金
       - 前 50%（含）→ 二等学业奖学金
       - > 50% → 不符合
    3. 排名缺失 → need_more_info

    【返回值字段说明】
    - decision:     "eligible" | "not_eligible" | "need_more_info"
    - eligible:     True/False/None（need_more_info 时为 None）
    - level:        "first_class" | "second_class" | None
    - reasons:      判断原因列表（为什么符合/不符合）
    - missing_information: 缺失的信息（如排名未知）
    - profile:      识别的用户画像（方便前端展示"我们识别到了什么"）
    """

    reasons: List[str] = []
    missing: List[str] = []

    # ---- 一票否决项检查（按严重程度排序） ----
    # 设计要点：一旦命中任一否决项，立即返回结果，不再继续检查后续条件。
    # 这样能快速失败（fail-fast），也避免产生相互矛盾的理由。

    if profile.get("has_failed_course") or (profile.get("failed_courses") or 0) > 0:
        reasons.append(
            f"有 {profile.get('failed_courses', 1)} 门课程不及格，不符合申请条件。"
        )
        return _build_result("not_eligible", profile, reasons, missing)

    if profile.get("has_discipline"):
        reasons.append("有处分记录，不符合申请条件。")
        return _build_result("not_eligible", profile, reasons, missing)

    if profile.get("has_academic_misconduct"):
        reasons.append("存在学术不端记录，不符合申请条件。")
        return _build_result("not_eligible", profile, reasons, missing)

    if profile.get("has_fake_material"):
        reasons.append("存在材料造假记录，不符合申请条件。")
        return _build_result("not_eligible", profile, reasons, missing)

    if profile.get("tutor_negative"):
        reasons.append("导师评价不合格，不符合申请条件。")
        return _build_result("not_eligible", profile, reasons, missing)

    # ---- 排名判断 ----
    # 排名是奖学金等级划分的核心依据
    rank = profile.get("rank_percent")
    if rank is None:
        missing.append(
            "成绩排名百分比未知，无法判断奖学金等级。"
            '请提供排名信息（如「前20%」）。'
        )
        return _build_result("need_more_info", profile, reasons, missing)

    if rank <= 20:
        reasons.append(f"成绩排名前 {rank}%，满足一等学业奖学金条件。")
        return _build_result("eligible", profile, reasons, missing, level="first_class")
    elif rank <= 50:
        reasons.append(f"成绩排名前 {rank}%，满足二等学业奖学金条件。")
        return _build_result("eligible", profile, reasons, missing, level="second_class")
    else:
        reasons.append(
            f"成绩排名前 {rank}%，超过奖学金申请排名上限（前 50%），不符合申请条件。"
        )
        return _build_result("not_eligible", profile, reasons, missing)


def _build_result(
    decision: str,
    profile: Dict[str, Any],
    reasons: List[str],
    missing: List[str],
    level: Optional[str] = None,
) -> Dict[str, Any]:
    """构建统一的判断结果字典。

    为什么单独一个函数？
    - 确保所有分支返回的字典结构完全一致。
    - 方便后续新增判断维度时，只需传不同的 reasons 即可。
    """
    return {
        "decision": decision,
        # eligible: 只有 decision=="eligible" 时才为 True；need_more_info 时为 None
        "eligible": (
            True if decision == "eligible"
            else (None if decision == "need_more_info" else False)
        ),
        "level": level,
        "reasons": reasons,
        "missing_information": missing,
        "profile": {
            "grade": profile.get("grade"),
            "rank_percent": profile.get("rank_percent"),
            "failed_courses": profile.get("failed_courses", 0),
            "has_failed_course": profile.get("has_failed_course", False),
            "has_discipline": profile.get("has_discipline", False),
            "has_academic_misconduct": profile.get("has_academic_misconduct", False),
            "has_fake_material": profile.get("has_fake_material", False),
            "tutor_negative": profile.get("tutor_negative", False),
        },
    }


# =========================================================
# 功能 2：材料清单 Word 导出
# =========================================================
# 使用 python-docx 生成 .docx 文件。选择 docx 而非 PDF 的原因：
# 1. python-docx 是纯 Python 库，无需系统依赖（不像 reportlab/fpdf）
# 2. 学生通常需要在清单上做笔记、打勾，Word 比 PDF 更方便编辑
# 3. Windows 上 Word 是默认安装的，兼容性好


def generate_checklist_docx(
    task_name: str,
    materials: List[str],
    steps: List[str],
    notes: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """使用 python-docx 生成 Word 材料清单文件。

    【文件命名策略】
    使用英文 + 时间戳命名（如 checklist_20260507_153012.docx），原因：
    - 中文文件名在 Windows URL 编码时可能出问题
    - 时间戳保证文件名不冲突，支持多次生成

    【存储路径】
    backend/app/storage/generated_files/
    与 uploaded_files/ 分开，便于区分"用户上传的"和"系统生成的"。

    【Word 文档结构】
    1. 标题（居中、18pt 加粗）
    2. 一、所需材料（编号列表）
    3. 二、办理步骤（编号列表）
    4. 三、注意事项（如有）
    5. 生成时间水印（灰色小字）

    返回 dict 包含 filename、file_path、download_url。
    """

    # ---- 延迟导入：python-docx 只在需要生成 Word 时才加载 ----
    # 这样做的好处：如果只是做资格判断（不需要生成 Word），就不会加载 docx 库
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    notes = notes or []

    # ---- 确保目录存在（幂等操作） ----
    generated_dir = settings.storage_dir / "generated_files"
    generated_dir.mkdir(parents=True, exist_ok=True)

    # ---- 生成文件名 ----
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"checklist_{timestamp}.docx"
    file_path = generated_dir / filename

    # ---- 构建 Word 文档 ----
    doc = Document()

    # 标题：居中、大号、加粗
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(task_name or "办事材料清单")
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    doc.add_paragraph()  # 空行分隔

    # 一、所需材料
    doc.add_heading("一、所需材料", level=2)
    for i, m in enumerate(materials, 1):
        doc.add_paragraph(f"{i}. {m}", style="List Number")

    # 二、办理步骤
    doc.add_heading("二、办理步骤", level=2)
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")

    # 三、注意事项（可选）
    if notes:
        doc.add_heading("三、注意事项", level=2)
        for n in notes:
            doc.add_paragraph(f"- {n}")

    # 底部水印：生成时间
    doc.add_paragraph()
    generated_time = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
    footer_para = doc.add_paragraph(
        f"（本文档由系统自动生成，生成时间：{generated_time}）"
    )
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(str(file_path))
    logger.info(f"[business_tools] checklist docx saved: {file_path}")

    # ---- 构建下载 URL ----
    # /api/files/download/{filename} 由 backend/app/api/files.py 提供
    download_url = f"/api/files/download/{filename}"

    return {
        "filename": filename,
        "file_path": str(file_path),
        "download_url": download_url,
    }
